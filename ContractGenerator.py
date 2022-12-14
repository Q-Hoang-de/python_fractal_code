import docx2pdf
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx2pdf import convert
import os


class ContractGenerator:
    def __init__(self, employeeName, employeeAddress, beginDate, title, salary, salaryInWord, hourWage,
                 terminationPeriod, probation):
        self.employeeName = employeeName
        self.employeeAddress = employeeAddress
        self.beginDate = beginDate
        self.title = title
        self.salary = salary
        self.salaryInWord = salaryInWord
        self.hourWage = hourWage
        self.terminationPeriod = terminationPeriod
        self.probation = probation

    def createEmployeeFullInfo(self):
        return '%s, %s' % (self.employeeName, self.employeeAddress)

    def calculateWeeklyHour(self):
        salaryInNumber = float(str(self.salary).replace(',', '.'))
        hourWageInNumber = float(str(self.hourWage).replace(',', '.'))
        return round((salaryInNumber/hourWageInNumber/26*13), 2)

    def toPdf(self, file):
        convert(file)

    def generateContract(self):
        # some fixed variable
        employer = 'Thanh Hung Tran \nInh. das Restaurant KoKoNo, Schnaitheimer Str.9, 89520 Heidenheim'
        employer_title = '-- nachfolgend Arbeitgeber genannt --'
        employee_title = '-- nachfolgend Arbeitnehmer genannt --'

        employee = self.createEmployeeFullInfo()
        weeklyHours = self.calculateWeeklyHour()

        # create actual document
        document = Document()
        document.add_heading('Arbeitsvertrag', 0)

        document.add_paragraph("Zwischen")
        document.add_paragraph().add_run(employer).bold = True
        document.add_paragraph(employer_title)

        document.add_paragraph().add_run().add_break(WD_BREAK.LINE_CLEAR_RIGHT)
        document.add_paragraph("und")
        document.add_paragraph().add_run().add_break(WD_BREAK.LINE_CLEAR_RIGHT)

        document.add_paragraph().add_run(employee).bold = True
        document.add_paragraph(employee_title)

        document.add_paragraph().add_run().add_break(WD_BREAK.LINE_CLEAR_RIGHT)

        document.add_paragraph("wird folgender Anstellungsvertrag geschlossen:")
        document.add_paragraph().add_run().add_break(WD_BREAK.LINE_CLEAR_RIGHT)

        ###
        firstSectionHeading = document.add_heading('?? 1 Arbeitsbeginn, T??tigkeitsbereich')
        document.add_paragraph('Der Arbeitnehmer tritt am %s.' % (self.beginDate))
        document.add_paragraph('Der Arbeitnehmer wird als %s besch??ftigt. ' \
                               'Der Arbeitgeber kann den dem Arbeitnehmer zugewiesenen Aufgabenbereich je nach den ' \
                               'gesch??ftlichen Erfordernissen erg??nzen oder auch ??ndern. Der Arbeitnehmer verpflichtet ' \
                               'sich dar??ber hinaus, vor??bergehend auch in anderen Betriebsst??tten des Arbeitgebers t??tig zu sein. ' \
                               'Der Anspruch des Arbeitnehmers auf die Gehaltszahlung nach Ma??gabe des ?? 3 dieses Vertrages bleibt ' \
                               'hiervon unber??hrt.' % (
                                   self.title)).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        ###
        secondSectionHeading = document.add_heading('?? 2 Arbeitszeit')
        secondSection = document.add_paragraph(
            'Die regelm????ige w??chentliche Arbeitszeit betr??gt %s Stunden.' \
            'Die Lage der Arbeitszeit und der Pausen richtet sich nach den betrieblichen Gepflogenheiten.' \
            % (str(weeklyHours)))
        ###
        thirdSectionHeading = document.add_heading('?? 3 Verg??tung')
        document.add_paragraph('F??r seine T??tigkeit erh??lt der Arbeitnehmer ein Monatsbruttogehalt in H??he von %s ???'
                               '\n(in Worten: %s).' % (str(self.salary), self.salaryInWord))
        document.add_paragraph('Stundenlohn betr??gt: %s Euro.' % (str(self.hourWage)))
        document.add_paragraph('Mit dem Gehalt sind s??mtliche Anspr??che des Arbeitnehmers auf ??berstunden ' \
                               'bzw. Mehrarbeit bzw. Sonn- und Feiertagsarbeit abgegolten. ' \
                               'Eine Verg??tung solcher Zeiten findet im ??brigen nur statt, wenn dies im Einzelfall vom ' \
                               'Arbeitgeber schriftlich zugesagt worden ist.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        ###
        fourthSectionHeading = document.add_heading('?? 4 Erholungsurlaub')
        document.add_paragraph('Der Urlaub wird nach dem Bundesurlaubsgesetzes geregelt.')
        document.add_paragraph(
            '??ber den gesetzlichen Urlaubsanspruch nach Absatz 1 hinaus hat der Arbeitnehmer einen ??bergesetzlichen Anspruch auf bezahlten Jahresurlaub von 0 '
            'weiteren Tagen. Der Arbeitnehmer hat daher einen Urlaubsanspruch von insgesamt 24 Arbeitstagen j??hrlich.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph(
            'Mit der Urlaubserteilung durch den Arbeitgeber erf??llt dieser zun??chst den Anspruch des Arbeitnehmers auf Urlaub nach Absatz 1, '
            'im Anschluss daran den etwaig weitergehenden Anspruch nach Absatz 2.'
            'Zeitpunkt und Dauer des Urlaubs richten sich nach den betrieblichen Notwendigkeiten und M??glichkeiten unter Ber??cksichtigung '
            'der W??nsche des Arbeitnehmers').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph('Das Urlaubsjahr ist das Kalenderjahr.')

        document.add_paragraph(
            'Der Urlaub muss im laufenden Kalenderjahr gew??hrt und genommen werden. Eine ??bertragung des Urlaubs auf das n??chste Kalenderjahr ist nur statthaft, '
            'wenn dringende betriebliche oder in der Person des Arbeitnehmers liegende Gr??nde dies rechtfertigen. '
            'Im Fall der ??bertragung muss der Urlaub innerhalb der ersten drei Monate des folgenden Kalenderjahres genommen werden. '
            'Andernfalls verf??llt er mit Ablauf des 31.03. des Folgejahres, soweit nicht zwingende gesetzliche Vorgaben etwas Anderes bestimmen.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph(
            'Kann der gesetzliche Urlaub wegen der Beendigung des Arbeitsverh??ltnisses ganz oder teilweise nicht gew??hrt werden, ist er abzugelten. '
            'In Bezug auf den gesetzlichen Urlaubsanspruch besteht ein Abgeltungsanspruch auch dann, wenn die Inanspruchnahme wegen krankheitsbedingter '
            'Arbeitsunf??higkeit nicht bis zum Ende des Kalenderjahres bzw. ??? f??r den Fall der ??bertragung ??? bis zum 31.03 des Folgejahres erfolgt ist. '
            'Dies gilt allerdings l??ngstenfalls bis zum 31.03. des ??bern??chsten Jahres. (Beispiel: Der Urlaubsanspruch f??r das Jahr 2016 verf??llt auch in '
            'F??llen krankheitsbedingter Arbeitsunf??higkeit sp??testens am 31.03.2018). ').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        document.add_paragraph('Eine Abgeltung des ??bergesetzlichen Urlaubsanspruchs ist ausgeschlossen. '
                               'Dieser Anspruch erlischt mit der Beendigung des Arbeitsverh??ltnisses ersatzlos. Anspr??che des Arbeitnehmers auf Urlaub sind nicht vererblich, '
                               'soweit sie nicht vor dem Tod des Arbeitnehmers entstanden sind. Anspr??che auf Urlaub und Urlaubsabgeltung, die den ??bergesetzlichen Urlaub betreffen, '
                               'sind in keinem Fall vererblich.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph(
            'F??r die Zeiten der Elternzeit gew??hrt der Arbeitgeber keinen (anteiligen) Urlaub. Ein Anspruch auf Urlaubsabgeltung besteht insoweit nicht. '
            'Der Arbeitgeber macht bereits jetzt von seinem Recht nach ?? 17 Abs. 1 Satz 1 BEEG Gebrauch.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        document.add_paragraph(
            'Im ??brigen gelten erg??nzend die Bestimmungen des Bundesurlaubsgesetzes in der jeweils geltenden Fassung.')
        ###
        fifthSectionHeading = document.add_heading('?? 5 Arbeitsverhinderung')
        document.add_paragraph('Der Arbeitnehmer verpflichtet sich, jede Arbeitsverhinderung unverz??glich, '
                               'tunlichst noch vor Dienstbeginn, dem Arbeitgeber unter Benennung der voraussichtlichen Verhinderungsdauer, '
                               'ggf. telefonisch, mitzuteilen.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph(
            'Im Krankheitsfall hat der Arbeitnehmer unverz??glich, sp??testens jedoch vor Ablauf des dritten Kalendertages, '
            'dem Arbeitgeber eine ??rztlich erstellte Arbeitsunf??higkeitsbescheinigung vorzulegen, aus der sich die voraussichtliche Dauer der '
            'Krankheit ergibt. Dauert die Krankheit l??nger an als in der ??rztlich erstellten Bescheinigung angegeben, so ist der '
            'Arbeitnehmer gleichfalls zur unverz??glichen Mitteilung und Vorlage einer weiteren Bescheinigung verpflichtet.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph(
            'Im Falle der Freistellung des Arbeitnehmers zur Pflege seines erkrankten Kindes erfolgt keine Entgeltfortzahlung.')
        document.add_paragraph(
            'Im ??brigen gelten f??r den Krankheitsfall die jeweils ma??geblichen gesetzlichen Bestimmungen.')

        ###
        sixthSectionHeading = document.add_heading('?? 6 Einstellungsfragebogen')
        document.add_paragraph('Der als Anlage beigef??gte Einstellungsfragebogen ist Bestandteil dieses Vertrages. '
                               'Der Arbeitnehmer versichert die Vollst??ndigkeit und Richtigkeit der gemachten Angaben').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        ###
        seventhSectionHeading = document.add_heading('?? 7 Nebenbesch??ftigung/Wettbewerbsverbot')
        document.add_paragraph(
            'Der Arbeitnehmer hat seine gesamte Arbeitskraft ausschlie??lich dem Arbeitgeber zur Verf??gung zu stellen. '
            'Eine Nebenbesch??ftigung w??hrend des Arbeitsverh??ltnisses darf nur mit vorheriger schriftlicher Zustimmung des Arbeitgebers ??bernommen werden.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        document.add_paragraph(
            'Der Arbeitnehmer verpflichtet sich, nach seinem Ausscheiden nicht aktiv um Kunden des Arbeitgebers zu werben.')

        ###
        document.add_heading('?? 8 K??ndigungsfristen')
        document.add_paragraph(
            'Das Arbeitsverh??ltnis wird %s eingegangen. Die ersten sechs Monate, also die Zeit bis zum %s gelten als Probezeit. '
            'W??hrend dieser Zeit kann das Arbeitsverh??ltnis von beiden Seiten mit einer Frist von zwei Wochen (?? 622 Abs. (3) BGB) gek??ndigt werden.' % (
                self.terminationPeriod, self.probation)).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        document.add_paragraph(
            'Nach Ablauf der Probezeit gelten die gesetzlichen K??ndigungsfristen. Verl??ngerte K??ndigungsfristen aufgrund '
            'verl??ngerter Betriebszugeh??rigkeiten gelten f??r beide Vertragsparteien.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        document.add_paragraph('Jede K??ndigung hat schriftlich zu erfolgen.')
        document.add_paragraph('Der Arbeitgeber ist berechtigt, den Arbeitnehmer nach Ausspruch einer K??ndigung '
                               'unter Fortzahlung der Verg??tung und Anrechnung auf Resturlaubsanspr??che von der Arbeitsleistung freizustellen.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        ###
        document.add_heading('?? 9 K??ndigung vor Dienstantritt/Vertragsstrafe')
        document.add_paragraph(
            'Eine K??ndigung vor Dienstantritt ist ausgeschlossen. K??ndigt der Arbeitnehmer vor Dienstantritt oder h??lt er im '
            'Falle einer K??ndigung die f??r ihn geltende K??ndigungsfrist nicht ein, gilt eine Vertragsstrafe in H??he eines Bruttomonatsverdienstes '
            '(vgl. ?? 3 dieses Vertrages) als vereinbart. Weitergehende Anspr??che des Arbeitgebers auf Schadenersatz und/oder '
            'Unterlassung bleiben hiervon unber??hrt.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        ###
        document.add_heading('?? 10 Krankheiten, Behinderung')
        document.add_paragraph(
            'Der Arbeitnehmer versichert, dass er nach seiner Kenntnis derzeit an keiner Krankheit leidet, die ihn an der ordnungsgem????en '
            'Wahrnehmung seiner in diesem Vertrag bestehenden Pflichten hindert.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph(
            'Der Arbeitnehmer best??tigt ferner, dass er kein behinderter Mensch oder ein diesem Gleichgestellter im Sinne des SGB IX ist. '
            'Der Arbeitnehmer verpflichtet sich, dem Arbeitgeber unverz??glich Mitteilung zu machen, wenn er einen Antrag auf '
            'Anerkennung als behinderter Mensch bzw. Gleichgestellter stellt.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        ###
        document.add_heading('?? 11 Verfallklausel')
        document.add_paragraph(
            'S??mtliche Anspr??che aus dem Arbeitsverh??ltnis sind von beiden Vertragsparteien innerhalb einer Frist von sechs Monaten nach F??lligkeit der jeweils anderen Vertragspartei '
            'schriftlich gegen??ber geltend zu machen. Erfolgt diese Geltendmachung nicht, gelten die Anspr??che als verfallen.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph(
            'Werden die nach Abs. (1) rechtzeitig geltend gemachten Anspr??che von der Gegenseite abgelehnt oder erkl??rt sich die Gegenseite nicht innerhalb von einem '
            'Monat nach der Geltendmachung des Anspruches, so verf??llt dieser, wenn er nicht innerhalb von zwei Monaten nach der Ablehnung '
            'oder dem Fristablauf gerichtlich anh??ngig gemacht wird.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        ###
        document.add_heading('?? 12 Schlussbestimmungen')
        document.add_paragraph(
            'M??ndliche Nebenabreden sind nicht getroffen worden. ??nderungen und/oder Erg??nzungen dieser '
            'Vereinbarung bed??rfen zu ihrer Rechtswirksamkeit der Schriftform. Dies gilt auch f??r ein Abweichen vom Schriftformerfordernis selbst.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        document.add_paragraph(
            'Sollte eine Bestimmung dieses Vertrages unwirksam sein oder werden, so wird die Wirksamkeit der ??brigen Bestimmungen '
            'davon nicht ber??hrt. Anstelle der unwirksamen Bestimmung werden die Parteien eine solche Bestimmung treffen, '
            'die dem mit der unwirksamen Bestimmung beabsichtigten Zweck am n??chsten kommt. Dies gilt auch f??r die Ausf??llung eventueller Vertragsl??cken.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        ###
        ### Date, Location
        document.add_paragraph().add_run().add_break(WD_BREAK.LINE_CLEAR_RIGHT)
        document.add_paragraph('Ort, Datum ').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        document.add_paragraph(
            '.........................................................................................').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        # row[1].text = 'Ort, Datum'
        ### Signature
        document.add_paragraph().add_run().add_break(WD_BREAK.LINE_CLEAR_RIGHT)
        data = (('.........................................................................................',
                 '.........................................................................................'),
                ('Unterschrift Arbeitsnehmer', 'Unterschrift Arbeitsgeber'))

        tableForSignature = document.add_table(rows=2, cols=2)
        tableForSignature.alignment = WD_TABLE_ALIGNMENT.CENTER
        for field, fieldName in data:
            row = tableForSignature.add_row().cells
            row[0].text = field
            row[1].text = fieldName
        ###

        ## Formatting document
        for para in document.paragraphs:
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.first_line_indent = Inches(-0.25)

        document.styles['Normal']
        filename = self.getFileName(self.employeeName, self.salary)
        document.save(filename)
        self.toPdf(filename)
        self.deleteDocFile(filename)

    def getFileName(self, employeeName, salary):
        return 'Contracts/' + employeeName + '_' + str(salary) + '.docx'

    def deleteDocFile(self, file):
        os.remove(file)
